#!/usr/bin/env python3
"""Publish one checked Kanoonak draft inside the current Kanoonak project folder.

The CLI accepts one canonical case-folder leaf, never a workspace or delivery
root. It derives the exact existing ``المسودات`` directory beneath the
host-provided current workspace, never creates case structure, never touches an
issued-rulings folder, and never changes the checked text.
"""

from __future__ import annotations

import argparse
import hashlib
import importlib.util
import io
import json
import os
import re
import sys
import unicodedata
import zipfile
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any, Callable


ALLOWED_KINDS = {"حكم", "حكم-تمهيدي", "قرار"}
ALLOWED_ROLES = {"title", "transition", "substantive", "plain"}
DIGEST_RE = re.compile(r"[0-9a-f]{64}\Z")
EXPECTED_UNIVERSAL_VERSION = "2026-08-26.1"
LOCAL_PERSONA_ABSENT = "غير موجود"


class DraftCreationError(ValueError):
    """A fail-closed input, formatting, or publication error."""


class _ExclusiveWriteError(DraftCreationError):
    """An exclusive leaf write failed after creation became possible."""

    def __init__(self, message: str, *, path: Path):
        super().__init__(message)
        self.path = path


@dataclass(frozen=True)
class PublicationResult:
    docx: Path
    metadata: Path
    ruling_text_sha256: str
    artifact_sha256: str


@dataclass(frozen=True)
class CaseDraftBoundary:
    """One complete canonical case drafts directory under process CWD."""

    case_folder: str
    case_directory: Path
    drafts_directory: Path


def _load_workspace_module():
    """Load only the reviewed sibling workspace helper."""

    path = Path(__file__).with_name("manage_workspace.py")
    module_name = "_kanoonak_manage_workspace_2026_08_26_1"
    spec = importlib.util.spec_from_file_location(module_name, path)
    if spec is None or spec.loader is None:
        raise DraftCreationError("workspace helper is unavailable")
    module = importlib.util.module_from_spec(spec)
    sys.modules[module_name] = module
    try:
        spec.loader.exec_module(module)
    except Exception:
        sys.modules.pop(module_name, None)
        raise
    for name, description in (
        ("parse_case_leaf", "canonical case parser"),
        ("_target_entries", "normalized entry lookup"),
        ("_identity_from_mapping", "case identity parser"),
        ("_existing_summary_identity", "existing summary identity parser"),
        ("_project_identity", "case identity projection"),
        ("load_strict_json", "strict JSON parser"),
    ):
        if not callable(getattr(module, name, None)):
            raise DraftCreationError(f"workspace helper lacks {description}")
    for name in (
        "CANONICAL_CASE_FILES",
        "CANONICAL_CASE_DIRECTORIES",
        "DRAFTS_DIRECTORY",
    ):
        if not hasattr(module, name):
            raise DraftCreationError(f"workspace helper lacks required constant {name}")
    return module


def _require_case_leaf(case_folder: str) -> None:
    if (
        not isinstance(case_folder, str)
        or not case_folder
        or case_folder in {".", ".."}
        or "/" in case_folder
        or "\\" in case_folder
        or "\x00" in case_folder
        or unicodedata.normalize("NFC", case_folder) != case_folder
    ):
        raise DraftCreationError("case folder must be one canonical NFC leaf")


def _case_key_from_folder(case_folder: str) -> tuple[str, int, int]:
    _require_case_leaf(case_folder)
    try:
        case_type, members = _load_workspace_module().parse_case_leaf(case_folder)
    except ValueError as exc:
        raise DraftCreationError(
            "case folder does not match the canonical case grammar"
        ) from exc
    if not members:
        raise DraftCreationError("canonical case folder has no case member")
    number, judicial_year = members[0]
    return case_type, number, judicial_year


def _require_metadata_case(metadata: dict[str, Any], case_folder: str) -> None:
    expected_type, expected_number, expected_year = _case_key_from_folder(case_folder)
    case = metadata.get("case") if isinstance(metadata, dict) else None
    if not isinstance(case, dict):
        raise DraftCreationError("metadata case tuple is required")
    actual = (case.get("type"), case.get("number"), case.get("judicial_year"))
    expected = (expected_type, expected_number, expected_year)
    if (
        actual != expected
        or type(actual[1]) is not int
        or type(actual[2]) is not int
    ):
        raise DraftCreationError("metadata case tuple does not match the selected case folder")


def resolve_case_draft_boundary(
    case_folder: str, case_identity: dict[str, Any]
) -> CaseDraftBoundary:
    """Resolve one complete canonical case under the host-provided root."""

    _require_case_leaf(case_folder)
    root = Path.cwd()
    workspace = _load_workspace_module()
    try:
        workspace.parse_case_leaf(case_folder)
        if not root.is_dir():
            raise DraftCreationError("current workspace root is not a directory")
        root_entries = workspace._target_entries(root, (case_folder,))
        raw_case_folder = root_entries.get(case_folder)
        if raw_case_folder is None:
            raise DraftCreationError("selected canonical case folder does not exist")
        case_directory = root / raw_case_folder
        if not case_directory.is_dir():
            raise DraftCreationError("selected canonical case folder is not a directory")
        required_names = (
            workspace.CANONICAL_CASE_FILES
            + workspace.CANONICAL_CASE_DIRECTORIES
        )
        case_entries = workspace._target_entries(case_directory, required_names)
        for name in workspace.CANONICAL_CASE_FILES:
            raw_name = case_entries.get(name)
            if raw_name is None or not (case_directory / raw_name).is_file():
                raise DraftCreationError(f"validated case is missing required file {name}")
        for name in workspace.CANONICAL_CASE_DIRECTORIES:
            raw_name = case_entries.get(name)
            if raw_name is None or not (case_directory / raw_name).is_dir():
                raise DraftCreationError(
                    f"validated case is missing required directory {name}"
                )
        try:
            expected_identity = workspace._identity_from_mapping(
                case_identity, minimal=True
            )
            existing_identity = workspace._existing_summary_identity(
                case_directory / case_entries["الملخص.md"]
            )
        except (KeyError, TypeError, ValueError, OSError) as exc:
            raise DraftCreationError("selected case identity is unreadable") from exc
        if (
            workspace._project_identity(expected_identity) != case_folder
            or existing_identity != expected_identity
        ):
            raise DraftCreationError(
                "selected case identity does not match the existing case summary"
            )
        raw_drafts = case_entries[workspace.DRAFTS_DIRECTORY]
        return CaseDraftBoundary(
            case_folder=case_folder,
            case_directory=case_directory,
            drafts_directory=case_directory / raw_drafts,
        )
    except ValueError as exc:
        raise DraftCreationError(str(exc)) from exc
    except OSError as exc:
        raise DraftCreationError("current workspace boundary is unavailable") from exc


def _load_universal_module():
    """Load only the reviewed sibling checker, never an ambient module."""
    path = Path(__file__).with_name("check_ruling_universal.py")
    spec = importlib.util.spec_from_file_location(
        "_kanoonak_check_ruling_universal_2026_08_26_1", path
    )
    if spec is None or spec.loader is None:
        raise DraftCreationError("universal ruling checker is unavailable")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    if getattr(module, "UNIVERSAL_VERSION", None) != EXPECTED_UNIVERSAL_VERSION:
        raise DraftCreationError(
            f"universal ruling checker version must be {EXPECTED_UNIVERSAL_VERSION}"
        )
    for name in ("check_ruling_text", "parse_paragraphs"):
        if not callable(getattr(module, name, None)):
            raise DraftCreationError(f"universal ruling checker lacks required API {name}")
    return module


def _universal_api() -> tuple[Callable[[str], dict[str, Any]], Callable[[str], list[str]]]:
    module = _load_universal_module()
    return module.check_ruling_text, module.parse_paragraphs


def _read_bytes(path: Path) -> bytes:
    return path.read_bytes()


def _read_exact(path: Path, expected: bytes) -> None:
    if _read_bytes(path) != expected:
        raise DraftCreationError(f"published payload reread mismatch: {path.name}")


def _write_complete_exclusive(path: Path, payload: bytes) -> None:
    """Create one final leaf without overwrite; preserve created-or-uncertain output."""

    flags = os.O_WRONLY | os.O_CREAT | os.O_EXCL | getattr(os, "O_BINARY", 0)
    fd = os.open(path, flags, 0o600)
    try:
        offset = 0
        while offset < len(payload):
            written = os.write(fd, payload[offset:])
            if written <= 0:
                raise OSError("exclusive final-leaf write was short")
            offset += written
        os.fsync(fd)
        if os.fstat(fd).st_size != len(payload):
            raise OSError("exclusive final-leaf write was short")
    except Exception as exc:
        try:
            os.close(fd)
        except OSError:
            pass
        raise _ExclusiveWriteError(
            f"exclusive final-leaf write failed after creating or possibly creating {path.name}",
            path=path,
        ) from exc
    try:
        os.close(fd)
        _read_exact(path, payload)
    except Exception as exc:
        raise _ExclusiveWriteError(
            f"exclusive final-leaf verification failed after creating or possibly creating {path.name}",
            path=path,
        ) from exc


def _validate_formatting(
    formatting: dict[str, Any], paragraphs: list[str]
) -> list[dict[str, str | None]]:
    if not isinstance(formatting, dict) or set(formatting) != {"paragraphs"}:
        raise DraftCreationError("formatting JSON must contain exactly 'paragraphs'")
    items = formatting["paragraphs"]
    if not isinstance(items, list) or len(items) != len(paragraphs):
        raise DraftCreationError("formatting paragraph count must match checked text")
    result: list[dict[str, str | None]] = []
    for index, (item, paragraph) in enumerate(zip(items, paragraphs), start=1):
        if not isinstance(item, dict) or set(item) != {"role", "opening_phrase"}:
            raise DraftCreationError(f"formatting paragraph {index} has unknown or missing keys")
        role = item["role"]
        opening = item["opening_phrase"]
        if role not in ALLOWED_ROLES:
            raise DraftCreationError(f"formatting paragraph {index} has unsupported role")
        if role == "substantive":
            if not isinstance(opening, str) or not opening or not paragraph.startswith(opening):
                raise DraftCreationError(
                    f"formatting paragraph {index} needs an exact non-empty opening prefix"
                )
        elif opening is not None:
            raise DraftCreationError(
                f"formatting paragraph {index} opening_phrase must be null for role {role}"
            )
        result.append({"role": role, "opening_phrase": opening})
    return result


def _set_run_format(run, *, bold: bool) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{slot}"), "Arial")
    for tag in ("sz", "szCs"):
        element = rpr.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            rpr.append(element)
        element.set(qn("w:val"), "32")
    rtl = rpr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        rpr.append(rtl)
    rtl.set(qn("w:val"), "1")


def _set_paragraph_format(paragraph, role: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, Twips

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
        if role in {"title", "transition"}
        else WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Twips(360) if role in {"transition", "substantive"} else None
    ppr = paragraph._p.get_or_add_pPr()
    bidi = ppr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        ppr.append(bidi)
    bidi.set(qn("w:val"), "1")


def _docx_bytes(
    paragraphs: list[str], formatting: list[dict[str, str | None]]
) -> bytes:
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt
    except ImportError as exc:  # pragma: no cover - guarded by declared dev dependency
        raise DraftCreationError("the host Documents runtime must provide python-docx") from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(16)
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{slot}"), "Arial")
    for tag in ("sz", "szCs"):
        element = rpr.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            rpr.append(element)
        element.set(qn("w:val"), "32")

    for text, item in zip(paragraphs, formatting):
        role = str(item["role"])
        opening = item["opening_phrase"]
        paragraph = document.add_paragraph()
        _set_paragraph_format(paragraph, role)
        if role == "substantive":
            assert isinstance(opening, str)
            lead = paragraph.add_run(opening)
            _set_run_format(lead, bold=True)
            if remainder := text[len(opening):]:
                tail = paragraph.add_run(remainder)
                _set_run_format(tail, bold=False)
        else:
            run = paragraph.add_run(text)
            _set_run_format(run, bold=role in {"title", "transition"})

    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def _extract_docx_text(payload: bytes) -> str:
    from xml.etree import ElementTree as ET

    try:
        with zipfile.ZipFile(io.BytesIO(payload)) as archive:
            if archive.testzip() is not None:
                raise DraftCreationError("DOCX ZIP CRC verification failed")
            root = ET.fromstring(archive.read("word/document.xml"))
    except (KeyError, ValueError, ET.ParseError, zipfile.BadZipFile) as exc:
        raise DraftCreationError("DOCX structural reopen failed") from exc
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    return "\n\n".join(
        "".join(node.text or "" for node in paragraph.findall(".//w:t", ns))
        for paragraph in root.findall(".//w:body/w:p", ns)
    )


def _verify_docx_text(payload: bytes, expected: str) -> None:
    if _extract_docx_text(payload) != expected:
        raise DraftCreationError("DOCX XML extraction did not preserve the checked ruling")
    try:
        from docx import Document

        reopened = Document(io.BytesIO(payload))
    except Exception as exc:
        raise DraftCreationError("python-docx structural reopen failed") from exc
    if "\n\n".join(paragraph.text for paragraph in reopened.paragraphs) != expected:
        raise DraftCreationError("python-docx reopen did not preserve the checked ruling")


def _metadata_yaml(metadata: dict[str, Any], filename: str) -> bytes:
    required = {
        "kanoonak", "state", "case", "kind", "based_on", "directive_version",
        "local_persona_updated", "artifact_sha256", "ruling_text_sha256",
    }
    missing = required - set(metadata)
    if missing:
        raise DraftCreationError(f"metadata missing required fields: {sorted(missing)}")
    if metadata["kanoonak"] != "draft" or metadata["state"] != "مسودة":
        raise DraftCreationError("DOCX helper creates drafts only")
    if metadata["kind"] not in ALLOWED_KINDS:
        raise DraftCreationError("metadata kind is unsupported")
    persona_updated = metadata["local_persona_updated"]
    if persona_updated != LOCAL_PERSONA_ABSENT:
        if not isinstance(persona_updated, str):
            raise DraftCreationError("metadata local_persona_updated is invalid")
        try:
            parsed_persona_date = date.fromisoformat(persona_updated)
        except ValueError as exc:
            raise DraftCreationError(
                "metadata local_persona_updated is invalid"
            ) from exc
        if parsed_persona_date.isoformat() != persona_updated:
            raise DraftCreationError("metadata local_persona_updated is invalid")
    if not isinstance(metadata["case"], dict) or not metadata["case"]:
        raise DraftCreationError("metadata case tuple is required")
    for field in ("artifact_sha256", "ruling_text_sha256"):
        if not DIGEST_RE.fullmatch(str(metadata[field])):
            raise DraftCreationError(f"metadata {field} is invalid")
    fields = [
        ("kanoonak", metadata["kanoonak"]), ("state", metadata["state"]),
        ("case", metadata["case"]), ("kind", metadata["kind"]),
    ]
    if "subject" in metadata:
        fields.append(("subject", metadata["subject"]))
    fields.extend([
        ("based_on", metadata["based_on"]),
        ("directive_version", metadata["directive_version"]),
        ("local_persona_updated", metadata["local_persona_updated"]),
        ("ruling_text_sha256", metadata["ruling_text_sha256"]),
        ("artifact_sha256", metadata["artifact_sha256"]),
        ("artifact", filename),
    ])
    try:
        lines = [
            f"{key}: {json.dumps(value, ensure_ascii=False, separators=(',', ':'))}"
            for key, value in fields
        ]
    except (TypeError, ValueError) as exc:
        raise DraftCreationError("metadata contains a non-JSON value") from exc
    return ("---\n" + "\n".join(lines) + "\n---\n").encode("utf-8")


def _verify_metadata_payload(
    payload: bytes, *, filename: str, ruling_digest: str, artifact_digest: str
) -> None:
    try:
        text = payload.decode("utf-8", errors="strict")
        lines = text.splitlines()
        if len(lines) < 3 or lines[0] != "---" or lines[-1] != "---":
            raise ValueError("front-matter fence missing")
        parsed: dict[str, Any] = {}
        for line in lines[1:-1]:
            key, separator, raw = line.partition(": ")
            if not separator or not key or key in parsed:
                raise ValueError("invalid or duplicate sidecar field")
            parsed[key] = json.loads(raw)
    except (UnicodeError, ValueError, json.JSONDecodeError) as exc:
        raise DraftCreationError("metadata sidecar structural reopen failed") from exc
    expected = {
        "artifact": filename,
        "ruling_text_sha256": ruling_digest,
        "artifact_sha256": artifact_digest,
    }
    if any(parsed.get(key) != value for key, value in expected.items()):
        raise DraftCreationError("metadata sidecar digest or artifact binding mismatch")


def _partial_publication_error(
    *, preserved: tuple[Path, ...], blocker: Path | None = None
) -> DraftCreationError:
    message = (
        "publication stopped; preserved created or possibly created item(s): "
        + ", ".join(str(path) for path in preserved)
    )
    if blocker is not None:
        message += f"; completion blocker: {blocker}"
    return DraftCreationError(message)


def _publish_to_drafts(
    *,
    delivery_dir: Path,
    kind: str,
    ruling_text: str,
    expected_ruling_sha256: str,
    formatting: dict[str, Any],
    metadata: dict[str, Any],
) -> PublicationResult:
    if kind not in ALLOWED_KINDS:
        raise DraftCreationError("kind must be حكم, حكم-تمهيدي, or قرار")
    if not isinstance(ruling_text, str):
        raise DraftCreationError("ruling_text must be a string")
    if (
        not isinstance(expected_ruling_sha256, str)
        or not DIGEST_RE.fullmatch(expected_ruling_sha256)
    ):
        raise DraftCreationError("expected_ruling_sha256 is invalid")
    if not isinstance(formatting, dict) or not isinstance(metadata, dict):
        raise DraftCreationError("formatting and metadata must be objects")
    delivery_dir = Path(os.path.abspath(os.fspath(delivery_dir)))
    if not delivery_dir.is_dir():
        raise DraftCreationError("publication directory must be an existing directory")

    check_ruling_text, parse_paragraphs = _universal_api()
    checked = check_ruling_text(ruling_text)
    if not isinstance(checked, dict) or checked.get("conforms") is not True:
        raise DraftCreationError("ruling text did not pass the universal checker")
    actual_ruling_sha256 = hashlib.sha256(ruling_text.encode("utf-8")).hexdigest()
    if checked.get("ruling_sha256") != actual_ruling_sha256:
        raise DraftCreationError("universal checker returned an inconsistent ruling digest")
    if actual_ruling_sha256 != expected_ruling_sha256:
        raise DraftCreationError(
            "checked ruling digest does not match expected_ruling_sha256"
        )
    paragraphs = parse_paragraphs(ruling_text)
    format_items = _validate_formatting(formatting, paragraphs)
    payload = _docx_bytes(paragraphs, format_items)
    _verify_docx_text(payload, ruling_text)
    artifact_sha256 = hashlib.sha256(payload).hexdigest()

    metadata = dict(metadata)
    if metadata.get("kind") != kind:
        raise DraftCreationError("metadata kind does not match the requested draft kind")
    metadata["kind"] = kind
    metadata["ruling_text_sha256"] = actual_ruling_sha256
    metadata["artifact_sha256"] = artifact_sha256

    for number in range(1, 100):
        basename = f"مسودة-{kind}-{number:02d}"
        target = delivery_dir / f"{basename}.docx"
        sidecar = delivery_dir / f"{basename}.metadata.yaml"

        # Either pre-existing member reserves the suffix. lexists also treats
        # a broken link as occupied, so no directory entry is overwritten.
        if os.path.lexists(target) or os.path.lexists(sidecar):
            continue

        metadata_payload = _metadata_yaml(metadata, target.name)
        _verify_metadata_payload(
            metadata_payload,
            filename=target.name,
            ruling_digest=actual_ruling_sha256,
            artifact_digest=artifact_sha256,
        )

        try:
            _write_complete_exclusive(sidecar, metadata_payload)
        except FileExistsError:
            # A race before this invocation exposed either member may use the
            # next suffix.
            continue
        except _ExclusiveWriteError as exc:
            raise _partial_publication_error(preserved=(exc.path,)) from exc

        try:
            _write_complete_exclusive(target, payload)
        except FileExistsError as exc:
            raise _partial_publication_error(
                preserved=(sidecar,), blocker=target
            ) from exc
        except OSError as exc:
            raise _partial_publication_error(
                preserved=(sidecar,), blocker=target
            ) from exc
        except _ExclusiveWriteError as exc:
            raise _partial_publication_error(
                preserved=(sidecar, exc.path)
            ) from exc
        except Exception as exc:
            raise _partial_publication_error(preserved=(sidecar, target)) from exc

        try:
            final_docx_bytes = _read_bytes(target)
            if hashlib.sha256(final_docx_bytes).hexdigest() != artifact_sha256:
                raise DraftCreationError("published DOCX digest mismatch")
            _verify_docx_text(final_docx_bytes, ruling_text)
            final_metadata_bytes = _read_bytes(sidecar)
            _read_exact(sidecar, metadata_payload)
            _verify_metadata_payload(
                final_metadata_bytes,
                filename=target.name,
                ruling_digest=actual_ruling_sha256,
                artifact_digest=artifact_sha256,
            )
        except Exception as exc:
            raise _partial_publication_error(
                preserved=(sidecar, target)
            ) from exc

        return PublicationResult(
            docx=target,
            metadata=sidecar,
            ruling_text_sha256=actual_ruling_sha256,
            artifact_sha256=artifact_sha256,
        )
    raise DraftCreationError("all two-digit draft numbers are already in use")



def create_workspace_draft(
    *,
    case_folder: str,
    case_identity: dict[str, Any],
    kind: str,
    ruling_text: str,
    expected_ruling_sha256: str,
    formatting: dict[str, Any],
    metadata: dict[str, Any],
) -> PublicationResult:
    """Publish beneath one complete canonical case in the given process root."""

    boundary = resolve_case_draft_boundary(case_folder, case_identity)
    _require_metadata_case(metadata, case_folder)
    return _publish_to_drafts(
        delivery_dir=boundary.drafts_directory,
        kind=kind,
        ruling_text=ruling_text,
        expected_ruling_sha256=expected_ruling_sha256,
        formatting=formatting,
        metadata=metadata,
    )


def _parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--case-folder", required=True)
    parser.add_argument("--case-identity-json", type=Path, required=True)
    parser.add_argument("--kind", choices=sorted(ALLOWED_KINDS), required=True)
    parser.add_argument("--ruling-file", type=Path, required=True)
    parser.add_argument("--expected-ruling-sha256", required=True)
    parser.add_argument("--formatting-json", type=Path, required=True)
    parser.add_argument("--metadata-json", type=Path, required=True)
    return parser


def main(argv: list[str] | None = None) -> int:
    args = _parser().parse_args(argv)
    try:
        workspace = _load_workspace_module()
        ruling_text = args.ruling_file.read_bytes().decode("utf-8", errors="strict")
        case_identity = workspace.load_strict_json(
            args.case_identity_json.read_text(encoding="utf-8")
        )
        metadata = workspace.load_strict_json(
            args.metadata_json.read_text(encoding="utf-8")
        )
        formatting = workspace.load_strict_json(
            args.formatting_json.read_text(encoding="utf-8")
        )
        if not all(
            isinstance(value, dict)
            for value in (case_identity, metadata, formatting)
        ):
            raise DraftCreationError(
                "case identity, metadata, and formatting JSON must be objects"
            )
        result = create_workspace_draft(
            case_folder=args.case_folder,
            case_identity=case_identity,
            kind=args.kind,
            ruling_text=ruling_text,
            expected_ruling_sha256=args.expected_ruling_sha256,
            formatting=formatting,
            metadata=metadata,
        )
    except (OSError, ValueError, UnicodeError, json.JSONDecodeError, DraftCreationError) as exc:
        print(f"DOCX draft creation refused: {exc}", file=sys.stderr)
        return 2
    print(json.dumps({
        "docx": str(result.docx), "metadata": str(result.metadata),
        "ruling_text_sha256": result.ruling_text_sha256,
        "artifact_sha256": result.artifact_sha256,
    }, ensure_ascii=True))
    return 0


if __name__ == "__main__":  # pragma: no cover
    raise SystemExit(main())
